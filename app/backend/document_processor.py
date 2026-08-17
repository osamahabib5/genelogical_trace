"""
Document processing utilities for extracting text and footnotes from various formats.
Extracts footnote references directly from Word XML to correctly link citations to text.
"""

import os
import logging
import zipfile
import re
from typing import List, Dict, Tuple
from pathlib import Path
import json

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from lxml import etree
except ImportError:
    etree = None

logger = logging.getLogger(__name__)

# Word XML namespace
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Regexes used by the direct (non-LLM) entity extraction pipeline
CAPS_NAME_RE = re.compile(r'\b(?:[A-Z]{2,}[\s“”"\'\u2019]*){2,}[A-Z]{2,}\b')
TITLE_NAME_RE = re.compile(r'\b(?:[A-Z][a-z]{1,15}\.?\s){1,2}[A-Z][a-z]{1,15}\b')
LOCATION_RE = re.compile(
    r'\b(?:in|at|of|to|near)\s+'
    r'((?:[A-Z][a-zA-Z]{1,30}\s?){0,2}'
    r'(?:County|Parish|City|Town|Township|Island|River|Bay|Creek)'
    r'|[A-Z][a-zA-Z\s]{2,40}?'
    r'(?:Virginia|Carolina|Tennessee|Texas|Louisiana|Iowa|Georgia|Maryland|Delaware|Pennsylvania|Ohio|Kentucky|Mississippi|Alabama|Missouri|Kansas|England|America|Africa|Jamaica|Barbados))'
)

# Words that should never start a person-name candidate
NAME_STOPWORDS = {
    'after', 'african', 'africans', 'although', 'however', 'unlike', 'when',
    'while', 'during', 'before', 'in', 'at', 'as', 'but', 'the', 'a', 'an',
    'for', 'from', 'to', 'with', 'without', 'their', 'his', 'her', 'it',
    'he', 'she', 'they', 'we', 'our', 'this', 'that', 'these', 'those',
    'there', 'then', 'if', 'and', 'or', 'on', 'off', 'of', 'about', 'above',
    'below', 'between', 'among', 'through', 'upon', 'within', 'beyond',
    'court', 'vital', 'first', 'second', 'third', 'documenting', 'journal',
    'volume', 'inside', 'editorial', 'products', 'further', 'introduction',
    'abstract', 'conclusion', 'figure', 'table', 'appendix', 'chapter',
    'section', 'part', 'new', 'north', 'south', 'east', 'west', 'old',
    'young', 'free', 'white', 'black', 'indian', 'mulatto', 'english',
    'american', 'native', 'accomack', 'king', 'queen', 'prince', 'princess',
    'general', 'governor', 'captain', 'president', 'senator', 'mister',
    'mrs', 'mr', 'dr', 'professor', 'mrs.', 'mr.', 'dr.',
}

# Words that disqualify a candidate as a person name
NAME_LOCATION_WORDS = {
    'county', 'parish', 'city', 'town', 'township', 'river', 'bay', 'creek',
    'island', 'journal', 'society',
}


class DocumentProcessor:
    """Process various document formats and extract text and footnotes"""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.json'}
    CHUNK_SIZE = 1000
    OVERLAP = 100

    @staticmethod
    def process_document(file_path: str) -> Tuple[str, List[str]]:
        """
        Process a document and return its full text and chunks.
        Returns: Tuple of (full_text, chunk_list)
        """
        file_ext = Path(file_path).suffix.lower()

        if file_ext not in DocumentProcessor.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_ext}")

        if file_ext == '.pdf':
            text = DocumentProcessor._extract_pdf(file_path)
        elif file_ext == '.docx':
            text = DocumentProcessor._extract_docx(file_path)
        elif file_ext == '.txt':
            text = DocumentProcessor._extract_txt(file_path)
        elif file_ext == '.json':
            text = DocumentProcessor._extract_json(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        chunks = DocumentProcessor._chunk_text(text)
        return text, chunks

    @staticmethod
    def extract_footnotes_from_docx(file_path: str) -> Dict[str, str]:
        """
        Extract footnotes from word/footnotes.xml.
        Returns dict mapping footnote ID to full citation text.
        e.g. {"1": "Gigi Best-Richardson is a Historian...", "6": "Accomack County..."}
        """
        if etree is None:
            logger.warning("lxml not installed — cannot extract footnotes")
            return {}

        footnotes = {}
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                if 'word/footnotes.xml' not in docx_zip.namelist():
                    logger.info("No footnotes.xml found in document")
                    return {}

                xml_content = docx_zip.read('word/footnotes.xml')
                tree = etree.fromstring(xml_content)

                for fn_el in tree.findall(f'.//{{{W}}}footnote'):
                    fn_id = fn_el.get(f'{{{W}}}id')
                    if fn_id is None or int(fn_id) <= 0:
                        continue

                    text_nodes = fn_el.findall(f'.//{{{W}}}t')
                    fn_text = ''.join((t.text or '') for t in text_nodes).strip()

                    if fn_text:
                        footnotes[fn_id] = fn_text

            logger.info(f"Extracted {len(footnotes)} footnotes from {Path(file_path).name}")
            return footnotes

        except Exception as e:
            logger.error(f"Error extracting footnotes: {e}")
            return {}

    @staticmethod
    def extract_paragraphs_with_footnote_refs(file_path: str) -> List[Dict]:
        """
        Parse word/document.xml and extract each paragraph as:
        {
            "text": "full paragraph text with [FN:6] markers",
            "footnote_refs": ["6", "7"]   # IDs of footnotes referenced in this paragraph
        }

        This correctly handles Word superscript footnote references which are
        stored as <w:footnoteReference w:id="6"/> elements in the XML,
        not as plain text characters.
        """
        if etree is None:
            logger.warning("lxml not installed")
            return []

        paragraphs = []
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                xml_content = docx_zip.read('word/document.xml')
                tree = etree.fromstring(xml_content)

                for para in tree.findall(f'.//{{{W}}}p'):
                    para_text_parts = []
                    footnote_refs = []

                    for child in para.iter():
                        tag = child.tag.replace(f'{{{W}}}', '')

                        if tag == 't' and child.text:
                            para_text_parts.append(child.text)

                        elif tag == 'footnoteReference':
                            fn_id = child.get(f'{{{W}}}id')
                            if fn_id and int(fn_id) > 0:
                                footnote_refs.append(fn_id)
                                # Insert a visible marker in the text so we can
                                # match footnotes back to chunks later
                                para_text_parts.append(f'[FN:{fn_id}]')

                    para_text = ''.join(para_text_parts).strip()
                    if para_text:
                        paragraphs.append({
                            "text": para_text,
                            "footnote_refs": footnote_refs
                        })

            logger.info(f"Extracted {len(paragraphs)} paragraphs with footnote refs")
            return paragraphs

        except Exception as e:
            logger.error(f"Error extracting paragraphs with footnote refs: {e}")
            return []

    @staticmethod
    def build_text_and_chunk_footnote_map(
        file_path: str,
        footnotes: Dict[str, str]
    ) -> Tuple[str, List[str], Dict[int, List[Dict]]]:
        """
        Main entry point for DOCX processing with footnotes.

        Returns:
            - full_text: complete document text with [FN:X] markers
            - chunks: list of text chunks
            - chunk_footnote_map: {chunk_index: [{"number": "6", "text": "Accomack..."}]}
        """
        paragraphs = DocumentProcessor.extract_paragraphs_with_footnote_refs(file_path)

        if not paragraphs:
            # Fall back to standard extraction
            text = DocumentProcessor._extract_docx(file_path)
            chunks = DocumentProcessor._chunk_text(text)
            return text, chunks, {}

        # Build full text from paragraphs (includes [FN:X] markers)
        full_text = '\n'.join(p['text'] for p in paragraphs)

        # Build chunks from full text
        chunks = DocumentProcessor._chunk_text(full_text)

        # Match footnote refs to chunks by scanning for [FN:X] markers
        chunk_footnote_map = {}
        fn_pattern = re.compile(r'\[FN:(\d+)\]')

        for chunk_idx, chunk_text in enumerate(chunks):
            matches = fn_pattern.findall(chunk_text)
            if matches:
                fn_list = []
                for fn_id in matches:
                    if fn_id in footnotes:
                        fn_list.append({
                            "number": fn_id,
                            "text": footnotes[fn_id]
                        })
                if fn_list:
                    chunk_footnote_map[chunk_idx] = fn_list

        # Clean [FN:X] markers from chunks before storing
        clean_chunks = [
            re.sub(r'\[FN:\d+\]', '', chunk).strip()
            for chunk in chunks
        ]
        clean_full_text = re.sub(r'\[FN:\d+\]', '', full_text).strip()

        logger.info(
            f"Built {len(clean_chunks)} chunks, "
            f"footnotes linked to {len(chunk_footnote_map)} chunks"
        )
        return clean_full_text, clean_chunks, chunk_footnote_map

    @staticmethod
    def match_footnotes_to_chunks(
        chunks: List[str],
        footnotes: Dict[str, str]
    ) -> Dict[int, List[Dict]]:
        """
        Legacy fallback: match footnotes by pattern matching in chunk text.
        Used when paragraphs with footnote refs are not available.
        """
        chunk_footnotes = {}
        ref_pattern = re.compile(r'\[FN:(\d+)\]|\[(\d+)\]|\((\d+)\)')

        for chunk_idx, chunk_text in enumerate(chunks):
            matches = ref_pattern.findall(chunk_text)
            found = []
            for match in matches:
                fn_num = match[0] or match[1] or match[2]
                if fn_num and fn_num in footnotes:
                    found.append({"number": fn_num, "text": footnotes[fn_num]})
            if found:
                chunk_footnotes[chunk_idx] = found

        return chunk_footnotes

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        try:
            import pypdf
            text = []
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    try:
                        text.append(page.extract_text())
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        if DocxDocument is None:
            raise ImportError("python-docx is not installed")
        try:
            doc = DocxDocument(file_path)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise

    @staticmethod
    def _extract_json(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"Error reading JSON: {e}")
            raise

    @staticmethod
    def _chunk_text(
        text: str,
        chunk_size: int = None,
        overlap: int = None
    ) -> List[str]:
        if chunk_size is None:
            chunk_size = DocumentProcessor.CHUNK_SIZE
        if overlap is None:
            overlap = DocumentProcessor.OVERLAP

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk.strip())
            start = end - overlap

        return [c for c in chunks if c]

    @staticmethod
    def extract_genealogical_entities(text: str) -> Dict:
        entities = {
            'names': [],
            'dates': [],
            'locations': [],
            'occupations': [],
            'relationships': []
        }

        # Dates: year spans, month-name spans, day-month-year, and fuzzy years
        date_patterns = [
            r'\b(?:(?:19|20)\d{2}|Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[^\n]*?\d{4}\b',
            r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December),?\s+\d{4}\b',
            r'\b(?:about|circa|c\.|before|after)\s+\d{4}\b',
        ]
        dates = set()
        for pat in date_patterns:
            dates.update(re.findall(pat, text, re.IGNORECASE))
        entities['dates'] = sorted(dates)

        # Relationships (keyword match)
        relationships = ['father', 'mother', 'son', 'daughter', 'brother', 'sister',
                        'husband', 'wife', 'widow', 'widower', 'grandfather', 'grandmother',
                        'aunt', 'uncle', 'cousin', 'parent', 'child', 'sibling', 'spouse', 'ancestor']
        entities['relationships'] = [
            rel for rel in relationships
            if re.search(rf'\b{rel}\b', text, re.IGNORECASE)
        ]

        # Occupations (keyword match)
        occupations = ['farmer', 'planter', 'doctor', 'teacher', 'merchant', 'soldier',
                       'laborer', 'carpenter', 'blacksmith', 'nurse', 'cook', 'servant',
                       'clergy', 'minister', 'attorney', 'sheriff', 'midwife',
                       'seamstress', 'overseer']
        entities['occupations'] = [
            occ for occ in occupations
            if re.search(rf'\b{occ}', text, re.IGNORECASE)
        ]

        # Names: ALL-CAPS genealogical entries, plus Title Case names on
        # lines/sentences that mention vital events (born, died, married...).
        names = set()
        sentences = re.split(r'(?<=[.;])\s+|\n+', text)
        for sentence in sentences:
            if not re.search(
                r'\b(?:born|died|married|buried|son of|daughter of|wife of|husband of|child of|widow of)\b',
                sentence, re.IGNORECASE
            ):
                continue
            names.update(m.strip() for m in CAPS_NAME_RE.findall(sentence))
            names.update(m.strip() for m in TITLE_NAME_RE.findall(sentence))

        def _is_valid_name(candidate: str) -> bool:
            candidate = candidate.strip('"“”\'’ .')
            if len(candidate) <= 3 or any(ch.isdigit() for ch in candidate):
                return False
            words = [w.lower().strip('.,"“”') for w in candidate.split() if w.strip('.,"“”')]
            if len(words) < 2:
                return False
            if words[0] in NAME_STOPWORDS:
                return False
            if any(w in NAME_LOCATION_WORDS for w in words):
                return False
            return True

        entities['names'] = sorted({n for n in names if _is_valid_name(n)})

        # Locations: place names after in/at/of/to/near, ending in
        # County/Parish/City/Town/Township... or a known state/country.
        entities['locations'] = sorted({
            m.group(1).strip()
            for m in LOCATION_RE.finditer(text)
            if len(m.group(1).strip()) > 2
        })

        return entities

    @staticmethod
    def extract_person_records(text: str, max_records: int = 100) -> List[Dict]:
        """
        Regex-based extraction of structured person records from text.
        No LLM involved — used by the direct upload pipeline.

        Looks for names, then scans sentences containing each name for
        born/died years and locations, plus occupation and relationship hints.
        """
        entities = DocumentProcessor.extract_genealogical_entities(text)
        names = entities['names']
        sentences = re.split(r'(?<=[.;])\s+|\n+', text)

        records = []
        seen = set()
        for name in names:
            key = name.lower()
            if key in seen:
                continue
            seen.add(key)

            rec = {
                'person_name': name,
                'birth_date': None,
                'birth_location': None,
                'death_date': None,
                'death_location': None,
                'occupation': None,
                'relation_type': None,
                'related_to': None
            }

            for sentence in sentences:
                if name.lower() not in sentence.lower():
                    continue
                lower = sentence.lower()

                if 'born' in lower:
                    born_seg = re.split(r'\bdied\b', sentence, flags=re.I)[0]
                    after_born = born_seg.lower().split('born', 1)[1] if 'born' in born_seg.lower() else ''
                    year = re.search(r'\b(1[5-9]\d{2})\b', after_born)
                    if year:
                        rec['birth_date'] = year.group(1)
                    loc = LOCATION_RE.search(born_seg)
                    if loc:
                        rec['birth_location'] = loc.group(1).strip()

                if 'died' in lower:
                    after_died = re.split(r'\bdied\b', sentence, flags=re.I)[-1]
                    year = re.search(r'\b(1[5-9]\d{2})\b', after_died)
                    if year:
                        rec['death_date'] = year.group(1)
                    loc = LOCATION_RE.search(after_died)
                    if loc:
                        rec['death_location'] = loc.group(1).strip()

                for occ in entities['occupations']:
                    if re.search(rf'\b{occ}\b', sentence, re.I):
                        rec['occupation'] = occ
                        break

                for rel in entities['relationships']:
                    m = re.search(
                        rf'\b{rel}\s+of\s+((?:[A-Z][a-zA-Z]{{1,20}}\s?){{1,3}}[A-Z][a-zA-Z]{{1,20}})',
                        sentence, re.I
                    )
                    if m:
                        rec['relation_type'] = rel
                        related = re.sub(r'\s+(?:and|the)$', '', m.group(1).strip(), flags=re.I)
                        related = re.sub(r'^the\s+', '', related, flags=re.I)
                        rec['related_to'] = related.strip()
                        break

                # Enough evidence — stop scanning sentences for this person
                if rec['birth_date'] or rec['death_date']:
                    break

            records.append(rec)
            if len(records) >= max_records:
                break

        return records